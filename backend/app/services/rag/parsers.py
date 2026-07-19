"""Format-specific document parsers.

Every parser produces a list of :class:`Segment` — text units annotated
with as much location detail as the format allows (page numbers for
PDFs, headings for DOCX/Markdown, line numbers for text and code, sheet
and row context for spreadsheets, tag paths for XML). Those locations
flow through chunking into citations, enabling navigation back to the
exact origin of an answer.

Parsers for formats with optional dependencies import them lazily and
raise :class:`DocumentProcessingError` with an actionable message when
missing.
"""

from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path

from app.core.exceptions import DocumentProcessingError, UnsupportedMediaError
from app.core.logging import get_logger

logger = get_logger(__name__)


@dataclass
class Segment:
    text: str
    page_number: int | None = None
    section: str | None = None
    line_start: int | None = None
    line_end: int | None = None


@dataclass
class ParsedDocument:
    segments: list[Segment]
    page_count: int | None = None
    metadata: dict = field(default_factory=dict)


CODE_EXTENSIONS = {
    ".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".c", ".cpp", ".h", ".hpp",
    ".cs", ".go", ".rs", ".rb", ".php", ".swift", ".kt", ".scala", ".sql",
    ".sh", ".ps1", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".r", ".m",
}

DOCUMENT_EXTENSIONS = {
    ".pdf", ".docx", ".txt", ".md", ".markdown", ".csv", ".tsv",
    ".xlsx", ".xls", ".xml", ".html", ".htm", ".json",
} | CODE_EXTENSIONS


def _read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise DocumentProcessingError(f"Could not decode '{path.name}' as text")


# ------------------------------------------------------------------- parsers

def _parse_pdf(path: Path) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentProcessingError(
            "PDF support requires the 'pypdf' package (pip install pypdf)"
        ) from exc

    reader = PdfReader(str(path))
    segments = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            segments.append(Segment(text=text, page_number=page_number))
    metadata = {}
    if reader.metadata:
        for key in ("title", "author", "subject"):
            value = getattr(reader.metadata, key, None)
            if value:
                metadata[key] = str(value)
    return ParsedDocument(segments=segments, page_count=len(reader.pages), metadata=metadata)


def _parse_docx(path: Path) -> ParsedDocument:
    try:
        import docx
    except ImportError as exc:
        raise DocumentProcessingError(
            "DOCX support requires the 'python-docx' package (pip install python-docx)"
        ) from exc

    document = docx.Document(str(path))
    segments = []
    current_heading: str | None = None
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name.startswith("Heading"):
            current_heading = text
        segments.append(Segment(text=text, section=current_heading))
    for table_index, table in enumerate(document.tables, start=1):
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells)
            for row in table.rows
            if any(cell.text.strip() for cell in row.cells)
        ]
        if rows:
            segments.append(
                Segment(text="\n".join(rows), section=f"Table {table_index}")
            )
    return ParsedDocument(segments=segments)


def _parse_plain_text(path: Path, *, track_headings: bool = False) -> ParsedDocument:
    """Text, Markdown and source code: line-accurate segments split on blank
    lines; Markdown headings become sections."""
    content = _read_text(path)
    lines = content.splitlines()
    segments: list[Segment] = []
    current_heading: str | None = None
    block: list[str] = []
    block_start = 1

    def flush(end_line: int) -> None:
        text = "\n".join(block).strip()
        if text:
            segments.append(
                Segment(
                    text=text,
                    section=current_heading,
                    line_start=block_start,
                    line_end=end_line,
                )
            )

    for number, line in enumerate(lines, start=1):
        heading = re.match(r"^(#{1,6})\s+(.*)", line) if track_headings else None
        if heading:
            flush(number - 1)
            block = []
            current_heading = heading.group(2).strip()
            block_start = number
            block.append(line)
        elif line.strip() == "":
            flush(number - 1)
            block = []
            block_start = number + 1
        else:
            if not block:
                block_start = number
            block.append(line)
    flush(len(lines))
    return ParsedDocument(segments=segments)


def _parse_csv(path: Path, delimiter: str = ",") -> ParsedDocument:
    content = _read_text(path)
    reader = csv.reader(io.StringIO(content), delimiter=delimiter)
    rows = [row for row in reader if any(cell.strip() for cell in row)]
    if not rows:
        return ParsedDocument(segments=[])

    header = rows[0]
    segments = []
    # Rows rendered as "column: value" records read naturally for retrieval.
    for index, row in enumerate(rows[1:], start=2):
        pairs = [
            f"{column.strip()}: {value.strip()}"
            for column, value in zip(header, row)
            if value.strip()
        ]
        if pairs:
            segments.append(
                Segment(text="; ".join(pairs), line_start=index, line_end=index)
            )
    return ParsedDocument(
        segments=segments, metadata={"columns": [column.strip() for column in header]}
    )


def _parse_excel(path: Path) -> ParsedDocument:
    try:
        import openpyxl
    except ImportError as exc:
        raise DocumentProcessingError(
            "Excel support requires the 'openpyxl' package (pip install openpyxl)"
        ) from exc

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    segments = []
    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        header = [str(cell) if cell is not None else "" for cell in rows[0]]
        for index, row in enumerate(rows[1:], start=2):
            pairs = [
                f"{column}: {cell}"
                for column, cell in zip(header, row)
                if cell is not None and str(cell).strip()
            ]
            if pairs:
                segments.append(
                    Segment(
                        text="; ".join(pairs),
                        section=sheet.title,
                        line_start=index,
                        line_end=index,
                    )
                )
    workbook.close()
    return ParsedDocument(segments=segments)


class _HTMLTextExtractor(HTMLParser):
    _SKIP = {"script", "style", "noscript"}

    def __init__(self) -> None:
        super().__init__()
        self.blocks: list[tuple[str, str | None]] = []  # (text, heading)
        self._heading: str | None = None
        self._skip_depth = 0
        self._capture_heading: str | None = None
        self._buffer: list[str] = []

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip_depth += 1
        elif re.fullmatch(r"h[1-6]", tag):
            self._capture_heading = ""

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skip_depth:
            self._skip_depth -= 1
        elif re.fullmatch(r"h[1-6]", tag) and self._capture_heading is not None:
            self._heading = self._capture_heading.strip() or self._heading
            self._capture_heading = None
        elif tag in {"p", "div", "li", "tr", "section", "article"}:
            self._flush()

    def handle_data(self, data):
        if self._skip_depth:
            return
        if self._capture_heading is not None:
            self._capture_heading += data
        elif data.strip():
            self._buffer.append(data.strip())

    def _flush(self):
        text = " ".join(self._buffer).strip()
        self._buffer = []
        if text:
            self.blocks.append((text, self._heading))

    def close(self):
        self._flush()
        super().close()


def _parse_html(path: Path) -> ParsedDocument:
    extractor = _HTMLTextExtractor()
    extractor.feed(_read_text(path))
    extractor.close()
    return ParsedDocument(
        segments=[Segment(text=text, section=heading) for text, heading in extractor.blocks]
    )


def _parse_xml(path: Path) -> ParsedDocument:
    import xml.etree.ElementTree as ET

    try:
        root = ET.fromstring(_read_text(path))
    except ET.ParseError as exc:
        raise DocumentProcessingError(f"Invalid XML: {exc}") from exc

    segments = []

    def walk(element: ET.Element, path_parts: list[str]) -> None:
        tag = re.sub(r"\{.*\}", "", element.tag)  # strip namespaces
        parts = path_parts + [tag]
        text = (element.text or "").strip()
        if text:
            segments.append(Segment(text=text, section="/".join(parts)))
        for child in element:
            walk(child, parts)

    walk(root, [])
    return ParsedDocument(segments=segments)


def _parse_json(path: Path) -> ParsedDocument:
    try:
        data = json.loads(_read_text(path))
    except json.JSONDecodeError as exc:
        raise DocumentProcessingError(f"Invalid JSON: {exc}") from exc

    segments = []

    def walk(node: object, key_path: str) -> None:
        if isinstance(node, dict):
            for key, value in node.items():
                walk(value, f"{key_path}.{key}" if key_path else str(key))
        elif isinstance(node, list):
            for index, value in enumerate(node):
                walk(value, f"{key_path}[{index}]")
        else:
            text = str(node).strip()
            if text:
                segments.append(Segment(text=f"{key_path}: {text}", section=key_path))

    walk(data, "")
    return ParsedDocument(segments=segments)


# ------------------------------------------------------------------ dispatch

def detect_format(filename: str) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in DOCUMENT_EXTENSIONS:
        raise UnsupportedMediaError(
            f"Unsupported document type '{extension or 'unknown'}'"
        )
    if extension in CODE_EXTENSIONS:
        return "code"
    return extension.lstrip(".")


def parse_document(path: Path, filename: str) -> ParsedDocument:
    extension = Path(filename).suffix.lower()
    try:
        if extension == ".pdf":
            return _parse_pdf(path)
        if extension == ".docx":
            return _parse_docx(path)
        if extension in {".md", ".markdown"}:
            return _parse_plain_text(path, track_headings=True)
        if extension == ".txt":
            return _parse_plain_text(path)
        if extension == ".csv":
            return _parse_csv(path)
        if extension == ".tsv":
            return _parse_csv(path, delimiter="\t")
        if extension in {".xlsx", ".xls"}:
            return _parse_excel(path)
        if extension in {".html", ".htm"}:
            return _parse_html(path)
        if extension == ".xml":
            return _parse_xml(path)
        if extension == ".json":
            return _parse_json(path)
        if extension in CODE_EXTENSIONS:
            return _parse_plain_text(path)
    except DocumentProcessingError:
        raise
    except Exception as exc:  # noqa: BLE001 — wrap any parser failure uniformly
        raise DocumentProcessingError(f"Failed to parse '{filename}': {exc}") from exc

    raise UnsupportedMediaError(f"Unsupported document type '{extension}'")
